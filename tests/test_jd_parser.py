"""Tests for src/data/jd_parser.py — parse_jd_docx()."""
import pytest

from src.data.jd_parser import parse_jd_docx

# ---------------------------------------------------------------------------
# Required top-level fields in the returned dict
# ---------------------------------------------------------------------------
REQUIRED_FIELDS = {
    "title",
    "company",
    "location",
    "experience_years",
    "must_have_skills",
    "nice_to_have_skills",
    "hard_disqualifiers",
    "preferred_locations",
    "salary_range_lpa",
    "full_text",
}


class TestFallbackBehaviour:
    """Tests for the fallback path (file not found)."""

    def test_fallback_jd_has_required_fields(self):
        """parse_jd_docx on a non-existent path returns a dict with all required keys."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        assert isinstance(result, dict), "Result should be a dict"
        missing = REQUIRED_FIELDS - result.keys()
        assert not missing, f"Missing fields in fallback result: {missing}"

    def test_must_have_skills_present(self):
        """must_have_skills is a non-empty list in the fallback."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        skills = result["must_have_skills"]
        assert isinstance(skills, list), "must_have_skills should be a list"
        assert len(skills) > 0, "must_have_skills should not be empty"

    def test_experience_years_present(self):
        """experience_years has both 'min' and 'max' keys in the fallback."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        exp = result["experience_years"]
        assert isinstance(exp, dict), "experience_years should be a dict"
        assert "min" in exp, "experience_years should have 'min'"
        assert "max" in exp, "experience_years should have 'max'"
        assert isinstance(exp["min"], int), "'min' should be an int"
        assert isinstance(exp["max"], int), "'max' should be an int"

    def test_preferred_locations(self):
        """preferred_locations contains Pune and Noida in the fallback."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        locs = result["preferred_locations"]
        assert isinstance(locs, list), "preferred_locations should be a list"
        assert "Pune" in locs, "preferred_locations should contain 'Pune'"
        assert "Noida" in locs, "preferred_locations should contain 'Noida'"

    def test_fallback_salary_range(self):
        """salary_range_lpa has min and max floats in the fallback."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        salary = result["salary_range_lpa"]
        assert isinstance(salary, dict)
        assert salary["min"] == 25.0
        assert salary["max"] == 55.0

    def test_fallback_must_have_includes_required_skills(self):
        """Fallback must_have_skills includes all four required skills."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        skills_lower = [s.lower() for s in result["must_have_skills"]]
        assert any("embedding" in s for s in skills_lower), \
            "embedding-based retrieval should be in must_have_skills"
        assert any("vector" in s for s in skills_lower), \
            "vector database should be in must_have_skills"
        assert any("python" in s for s in skills_lower), \
            "Python should be in must_have_skills"
        assert any("ndcg" in s or "evaluation" in s for s in skills_lower), \
            "evaluation frameworks should be in must_have_skills"

    def test_fallback_hard_disqualifiers_present(self):
        """Fallback hard_disqualifiers contains the three required entries."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        dq = [d.lower() for d in result["hard_disqualifiers"]]
        assert any("consulting" in d for d in dq), \
            "consulting-only career should be a hard disqualifier"
        assert any("llm" in d for d in dq), \
            "LLM-only experience should be a hard disqualifier"
        assert any("production" in d for d in dq), \
            "no production deployment should be a hard disqualifier"

    def test_fallback_full_text_is_string(self):
        """full_text in the fallback is a non-empty string."""
        result = parse_jd_docx("/nonexistent/path/job_description.docx")
        assert isinstance(result["full_text"], str)
        assert len(result["full_text"]) > 0

    def test_fallback_is_independent_copy(self):
        """Two calls return independent dicts (mutating one does not affect another)."""
        r1 = parse_jd_docx("/nonexistent/a.docx")
        r2 = parse_jd_docx("/nonexistent/b.docx")
        r1["must_have_skills"].append("MUTATED")
        assert "MUTATED" not in r2["must_have_skills"]


class TestDocxParsing:
    """Tests for actual docx file parsing."""

    def test_parse_real_docx_returns_required_fields(self, tmp_path):
        """parse_jd_docx on a real docx file returns all required fields."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_file = tmp_path / "test_jd.docx"
        doc = Document()
        doc.add_paragraph("Senior AI Engineer — Founding Team")
        doc.add_paragraph("Company: Redrob AI")
        doc.add_paragraph("Location: Pune/Noida, India (Hybrid)")
        doc.add_paragraph("5-9 years experience required")
        doc.add_paragraph("Salary: 25-55 LPA")
        doc.add_paragraph("Must Have:")
        doc.add_paragraph("- embedding-based retrieval")
        doc.add_paragraph("- vector database")
        doc.add_paragraph("- Python")
        doc.add_paragraph("- evaluation frameworks (NDCG, MRR, MAP)")
        doc.save(str(docx_file))

        result = parse_jd_docx(str(docx_file))
        missing = REQUIRED_FIELDS - result.keys()
        assert not missing, f"Missing fields: {missing}"

    def test_parse_real_docx_extracts_experience_range(self, tmp_path):
        """Parser extracts a valid experience range from a real docx."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_file = tmp_path / "exp_test.docx"
        doc = Document()
        doc.add_paragraph("Test JD")
        doc.add_paragraph("We need 5-9 years of experience in AI/ML systems.")
        doc.save(str(docx_file))

        result = parse_jd_docx(str(docx_file))
        exp = result["experience_years"]
        assert exp["min"] == 5
        assert exp["max"] == 9

    def test_parse_real_docx_extracts_preferred_locations(self, tmp_path):
        """Parser extracts Pune and Noida from a real docx location line."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_file = tmp_path / "loc_test.docx"
        doc = Document()
        doc.add_paragraph("Test JD")
        doc.add_paragraph("Location: Pune/Noida, India (Hybrid)")
        doc.save(str(docx_file))

        result = parse_jd_docx(str(docx_file))
        locs = result["preferred_locations"]
        assert "Pune" in locs
        assert "Noida" in locs
